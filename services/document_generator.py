"""
Service de génération de documents : txt, docx, pdf
"""
import os
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Génère les documents finaux : Minutes, Pré-CR, Relevé des décisions"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        # Enregistrement d'une police pour le PDF (optionnel)
        try:
            # Vous pouvez ajouter une police personnalisée ici
            pass
        except:
            pass
    
    def generate_all_documents(self, session_id: str, transcription: Dict[str, Any],
                              speaker_mapping: Dict[str, str], pre_cr: str,
                              decisions: List[Dict[str, Any]], date_seance: str,
                              output_folder: str) -> Dict[str, str]:
        """
        Génère tous les documents pour une session
        
        Returns:
            dict: Chemins des fichiers générés
        """
        try:
            logger.info(f"Génération des documents pour la session {session_id}")
            
            output_path = Path(output_folder) / session_id
            output_path.mkdir(parents=True, exist_ok=True)
            
            # Format de la date pour le préfixe des noms de fichiers
            date_prefix = date_seance.replace('-', '') if date_seance else datetime.now().strftime('%Y%m%d')
            
            documents = {}
            
            # 1. Minutes (transcription verbatim)
            minutes_txt = self._generate_minutes_txt(
                transcription, speaker_mapping, date_seance
            )
            minutes_path = output_path / f"{date_prefix}_Minutes.txt"
            with open(minutes_path, 'w', encoding='utf-8') as f:
                f.write(minutes_txt)
            documents['minutes_txt'] = str(minutes_path)
            
            minutes_docx = self._generate_minutes_docx(
                transcription, speaker_mapping, date_seance, minutes_path.with_suffix('.docx')
            )
            documents['minutes_docx'] = str(minutes_docx)
            
            minutes_pdf = self._generate_minutes_pdf(
                transcription, speaker_mapping, date_seance, minutes_path.with_suffix('.pdf')
            )
            documents['minutes_pdf'] = str(minutes_pdf)
            
            # 2. Pré-compte rendu
            pre_cr_txt = self._generate_pre_cr_txt(pre_cr, date_seance)
            pre_cr_path = output_path / f"{date_prefix}_Pre-Compte-rendu.txt"
            with open(pre_cr_path, 'w', encoding='utf-8') as f:
                f.write(pre_cr_txt)
            documents['pre_cr_txt'] = str(pre_cr_path)
            
            pre_cr_docx = self._generate_pre_cr_docx(
                pre_cr, date_seance, pre_cr_path.with_suffix('.docx')
            )
            documents['pre_cr_docx'] = str(pre_cr_docx)
            
            pre_cr_pdf = self._generate_pre_cr_pdf(
                pre_cr, date_seance, pre_cr_path.with_suffix('.pdf')
            )
            documents['pre_cr_pdf'] = str(pre_cr_pdf)
            
            # 3. Relevé des décisions
            decisions_txt = self._generate_decisions_txt(decisions, date_seance)
            decisions_path = output_path / f"{date_prefix}_Releve-des-decisions.txt"
            with open(decisions_path, 'w', encoding='utf-8') as f:
                f.write(decisions_txt)
            documents['decisions_txt'] = str(decisions_path)
            
            decisions_docx = self._generate_decisions_docx(
                decisions, date_seance, decisions_path.with_suffix('.docx')
            )
            documents['decisions_docx'] = str(decisions_docx)
            
            decisions_pdf = self._generate_decisions_pdf(
                decisions, date_seance, decisions_path.with_suffix('.pdf')
            )
            documents['decisions_pdf'] = str(decisions_pdf)
            
            logger.info(f"Tous les documents générés avec succès")
            return documents
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération des documents: {str(e)}", exc_info=True)
            raise
    
    def _generate_minutes_txt(self, transcription: Dict[str, Any],
                             speaker_mapping: Dict[str, str], date_seance: str) -> str:
        """Génère le fichier TXT des minutes"""
        lines = []
        lines.append("=" * 80)
        lines.append("MINUTES DE LA RÉUNION")
        lines.append("=" * 80)
        lines.append(f"Date de la séance: {date_seance}")
        lines.append("")
        lines.append("TRANSCRIPTION VERBATIM")
        lines.append("-" * 80)
        lines.append("")
        
        segments = transcription.get('segments', [])
        for seg in segments:
            speaker_label = seg.get('speaker', 'UNKNOWN')
            speaker_name = speaker_mapping.get(speaker_label, speaker_label)
            text = seg.get('text', '')
            start = seg.get('start', 0)
            
            time_str = self._format_time(start)
            lines.append(f"[{time_str}] {speaker_name}:")
            lines.append(f"{text}")
            lines.append("")
        
        return "\n".join(lines)
    
    def _generate_minutes_docx(self, transcription: Dict[str, Any],
                              speaker_mapping: Dict[str, str], date_seance: str,
                              output_path: Path) -> Path:
        """Génère le fichier DOCX des minutes"""
        doc = Document()
        
        # Titre
        title = doc.add_heading('MINUTES DE LA RÉUNION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Date de la séance: {date_seance}")
        doc.add_paragraph("")
        
        # Sous-titre
        subtitle = doc.add_heading('TRANSCRIPTION VERBATIM', level=1)
        
        # Contenu
        segments = transcription.get('segments', [])
        for seg in segments:
            speaker_label = seg.get('speaker', 'UNKNOWN')
            speaker_name = speaker_mapping.get(speaker_label, speaker_label)
            text = seg.get('text', '')
            start = seg.get('start', 0)
            
            time_str = self._format_time(start)
            p = doc.add_paragraph()
            p.add_run(f"[{time_str}] ").bold = True
            p.add_run(f"{speaker_name}:").bold = True
            doc.add_paragraph(text)
            doc.add_paragraph("")
        
        doc.save(str(output_path))
        return output_path
    
    def _generate_minutes_pdf(self, transcription: Dict[str, Any],
                             speaker_mapping: Dict[str, str], date_seance: str,
                             output_path: Path) -> Path:
        """Génère le fichier PDF des minutes"""
        doc = SimpleDocTemplate(str(output_path), pagesize=A4, title="Minutes de la réunion")
        story = []
        
        # Styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor='black',
            spaceAfter=30,
            alignment=1  # Centré
        )
        
        # Titre
        story.append(Paragraph("MINUTES DE LA RÉUNION", title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Date de la séance: {date_seance}", self.styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Contenu
        segments = transcription.get('segments', [])
        for seg in segments:
            speaker_label = seg.get('speaker', 'UNKNOWN')
            speaker_name = speaker_mapping.get(speaker_label, speaker_label)
            text = seg.get('text', '')
            start = seg.get('start', 0)
            
            time_str = self._format_time(start)
            speaker_text = f"<b>[{time_str}] {speaker_name}:</b>"
            story.append(Paragraph(speaker_text, self.styles['Normal']))
            story.append(Paragraph(text, self.styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
        return output_path
    
    def _generate_pre_cr_txt(self, pre_cr: str, date_seance: str) -> str:
        """Génère le fichier TXT du pré-compte rendu"""
        lines = []
        lines.append("=" * 80)
        lines.append("PRÉ-COMPTE RENDU DE LA RÉUNION")
        lines.append("=" * 80)
        lines.append(f"Date de la séance: {date_seance}")
        lines.append("")
        lines.append("-" * 80)
        lines.append("")
        lines.append(pre_cr)
        return "\n".join(lines)
    
    def _generate_pre_cr_docx(self, pre_cr: str, date_seance: str,
                             output_path: Path) -> Path:
        """Génère le fichier DOCX du pré-compte rendu"""
        doc = Document()
        
        title = doc.add_heading('PRÉ-COMPTE RENDU DE LA RÉUNION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date de la séance: {date_seance}")
        doc.add_paragraph("")
        
        # Ajout du contenu paragraphe par paragraphe
        for para in pre_cr.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(str(output_path))
        return output_path
    
    def _generate_pre_cr_pdf(self, pre_cr: str, date_seance: str,
                            output_path: Path) -> Path:
        """Génère le fichier PDF du pré-compte rendu"""
        doc = SimpleDocTemplate(str(output_path), pagesize=A4, title="Pré-compte rendu de la réunion")
        story = []
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor='black',
            spaceAfter=30,
            alignment=1
        )
        
        story.append(Paragraph("PRÉ-COMPTE RENDU DE LA RÉUNION", title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Date de la séance: {date_seance}", self.styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        for para in pre_cr.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), self.styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
        return output_path
    
    def _generate_decisions_txt(self, decisions: List[Dict[str, Any]], date_seance: str) -> str:
        """Génère le fichier TXT du relevé des décisions"""
        lines = []
        lines.append("=" * 80)
        lines.append("RELEVÉ DES DÉCISIONS")
        lines.append("=" * 80)
        lines.append(f"Date de la séance: {date_seance}")
        lines.append("")
        lines.append("-" * 80)
        lines.append("")
        
        if not decisions:
            lines.append("Aucune décision enregistrée.")
        else:
            for i, decision in enumerate(decisions, 1):
                lines.append(f"DÉCISION N° {decision.get('numero', i)}")
                lines.append("")
                if decision.get('titre'):
                    lines.append(f"Titre: {decision['titre']}")
                if decision.get('description'):
                    lines.append(f"Description: {decision['description']}")
                if decision.get('vote'):
                    lines.append(f"Vote: {decision['vote']}")
                if decision.get('timestamp'):
                    lines.append(f"Timestamp: {decision['timestamp']}")
                lines.append("")
                lines.append("-" * 80)
                lines.append("")
        
        return "\n".join(lines)
    
    def _generate_decisions_docx(self, decisions: List[Dict[str, Any]], date_seance: str,
                                output_path: Path) -> Path:
        """Génère le fichier DOCX du relevé des décisions"""
        doc = Document()
        
        title = doc.add_heading('RELEVÉ DES DÉCISIONS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date de la séance: {date_seance}")
        doc.add_paragraph("")
        
        if not decisions:
            doc.add_paragraph("Aucune décision enregistrée.")
        else:
            for i, decision in enumerate(decisions, 1):
                heading = doc.add_heading(f"DÉCISION N° {decision.get('numero', i)}", level=1)
                if decision.get('titre'):
                    p = doc.add_paragraph()
                    p.add_run("Titre: ").bold = True
                    p.add_run(decision['titre'])
                if decision.get('description'):
                    p = doc.add_paragraph()
                    p.add_run("Description: ").bold = True
                    p.add_run(decision['description'])
                if decision.get('vote'):
                    p = doc.add_paragraph()
                    p.add_run("Vote: ").bold = True
                    p.add_run(decision['vote'])
                if decision.get('timestamp'):
                    p = doc.add_paragraph()
                    p.add_run("Timestamp: ").bold = True
                    p.add_run(decision['timestamp'])
                doc.add_paragraph("")
        
        doc.save(str(output_path))
        return output_path
    
    def _generate_decisions_pdf(self, decisions: List[Dict[str, Any]], date_seance: str,
                               output_path: Path) -> Path:
        """Génère le fichier PDF du relevé des décisions"""
        doc = SimpleDocTemplate(str(output_path), pagesize=A4, title="Relevé des décisions")
        story = []
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor='black',
            spaceAfter=30,
            alignment=1
        )
        
        story.append(Paragraph("RELEVÉ DES DÉCISIONS", title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Date de la séance: {date_seance}", self.styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        if not decisions:
            story.append(Paragraph("Aucune décision enregistrée.", self.styles['Normal']))
        else:
            for i, decision in enumerate(decisions, 1):
                story.append(Paragraph(
                    f"DÉCISION N° {decision.get('numero', i)}",
                    self.styles['Heading2']
                ))
                story.append(Spacer(1, 0.1*inch))
                
                if decision.get('titre'):
                    story.append(Paragraph(
                        f"<b>Titre:</b> {decision['titre']}",
                        self.styles['Normal']
                    ))
                if decision.get('description'):
                    story.append(Paragraph(
                        f"<b>Description:</b> {decision['description']}",
                        self.styles['Normal']
                    ))
                if decision.get('vote'):
                    story.append(Paragraph(
                        f"<b>Vote:</b> {decision['vote']}",
                        self.styles['Normal']
                    ))
                if decision.get('timestamp'):
                    story.append(Paragraph(
                        f"<b>Timestamp:</b> {decision['timestamp']}",
                        self.styles['Normal']
                    ))
                
                story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        return output_path
    
    def _format_time(self, seconds: float) -> str:
        """Formate les secondes en HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

